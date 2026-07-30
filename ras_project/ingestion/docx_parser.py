"""ingestion/docx_parser.py — Parse DOCX files into page-like dicts."""
import os
from typing import Dict, Any
from utils.logger import get_logger
logger = get_logger(__name__)
_CPP = 3000  # chars per virtual page


class DOCXParser:
    def parse(self, file_path: str) -> Dict[str, Any]:
        logger.info(f"Parsing DOCX: {file_path}")
        pages, metadata = [], {}
        try:
            from docx import Document
            doc = Document(file_path)
            metadata = {
                "title":  doc.core_properties.title or "",
                "author": doc.core_properties.author or "",
            }
            full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            tables_data = []
            for tbl in doc.tables:
                tables_data.append([[c.text.strip() for c in row.cells] for row in tbl.rows])
            for i, chunk in enumerate(_split(full_text, _CPP), 1):
                pages.append({"page_number":i,"text":chunk,
                              "tables": tables_data if i==1 else [],
                              "source":os.path.basename(file_path),"is_scanned":False})
        except ImportError:
            logger.error("python-docx not installed.")
        except Exception as e:
            logger.error(f"DOCX parse error: {e}")
        if not pages:
            pages = [{"page_number":1,"text":"","tables":[],
                      "source":os.path.basename(file_path),"is_scanned":False}]
        return {"source": file_path, "metadata": metadata, "pages": pages}


def _split(text, n):
    return [text[i:i+n] for i in range(0, max(len(text),1), n)] or [""]
